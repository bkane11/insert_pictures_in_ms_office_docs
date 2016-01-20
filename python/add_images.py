from __future__ import print_function

import os
import shutil
# import xlsxwriter
import openpyxl

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import re
# import pandas as pd

# cwd = os.getcwd()

def addImageToSheet(worksheet, imagepath, cellno, text=None, insertOptions={}):
  img = openpyxl.drawing.image.Image( imagepath, size=(1000,1000) )
  # img = openpyxl.drawing.Image(os.path.join(cwd, imagename) )
  img.anchor(worksheet.cell('A'+str(cellno)))
  worksheet.add_image(img)
  return worksheet

def addWorkSheet(wb, name):
   return wb.create_sheet(title=name)

def update_xlsx(name=None, onesheet=True, images=[], then=None):
    filepath = os.path.realpath(name)
    # filepath = os.path.join(cwd, name)
    print('updating:', filepath)

    # images = reduce(lambda a,b: a+b, map(lambda x: x.split(','), images ) )
    # print(1, 'images', images)

    #Open an xlsx for reading
    wb = openpyxl.load_workbook(filename=filepath)
    #You can also select a particular sheet
    #based on sheet name
    storedsheet = None
    index = 0
    regex = re.compile(r".+\/(.+)\.(jpg|png|bmp)$", re.IGNORECASE)
    for image in images:
        index += 1      
        row = 1 if not onesheet else index*25 if index>1 else 1
        # print('image:', image)
        # return
        matches = regex.match(image)
        # print(matches)
        sheetname = matches.groups()[0] if not onesheet else 'Images'
        # sheetname = regex.sub('', image) if not onesheet else 'Images'
        
        # return print('imagename is', image, 'sheetname is:',sheetname)

        image = os.path.realpath(image)
        if not os.path.exists(image):
          print({'error': image + ' does not exist'})
          continue

        print('adding image:', image, 'to sheet', sheetname,' at index:', row)
        
        if not onesheet:
            ws = addWorkSheet(wb, sheetname)
            addImageToSheet(ws, image , row)
        else:
            storedsheet = storedsheet if storedsheet else addWorkSheet(wb, sheetname)
            addImageToSheet(storedsheet, image , row)
            # storedsheet = ws
        
        # addImageToSheet(ws, image, index*25)
        # wb.save(filepath)
    
    wb.save(filepath)
    
    if type(then).__name__ == 'function':
      then(filepath)

    return filepath

def update_docx(name=None, onesheet=True, images=[], then=None):
    filepath = os.path.realpath(name)
    # filepath = os.path.join(cwd, name)
    print('updating:', filepath)

    # images = reduce(lambda a,b: a+b, map(lambda x: x.split(','), images ) )
    print(1, 'images', images)

    #Open an docx for reading
    doc = docx.Document(docx=filepath)

    regex = re.compile(r".+\/(.+)\.(jpg|png|bmp)$", re.IGNORECASE)
    
    for image in images:
        imagename = regex.match(image).groups()[0]
        image = os.path.realpath(image)
        if not os.path.exists(image):
          print({'error': image + ' does not exist'})
          continue

        print('adding image:', image, 'to sheet', filepath)
        
        p = doc.add_paragraph()
        paragraph_format = p.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.font.bold = True
        r.font.size = Pt(16.0)
        r.add_text(imagename)
        # r.add_picture(image)
        r.add_picture(image, height=8500000)
        # doc.inline_shapes[len(doc.inline_shapes)-1].height = 4000000
        # print(doc.inline_shapes[len(doc.inline_shapes)-1].height )
    
    os.remove(filepath)
    doc.save(filepath)
    
    if type(then).__name__ == 'function':
      then(filepath)

    return filepath

testfile = './tests/inputs/test.docx'
# testfile = './tests/inputs/test.xlsx'
testimages = ['./tests/images/image_1.jpg', './tests/images/image_2.jpg']

def addImages(name=testfile, images=testimages, onesheet=False, then=None):
  if re.search('\.xlsx$', name):
    return update_xlsx(name=name, onesheet=onesheet, images=images, then=then)
  elif re.search('\.docx$', name):
    return update_docx(name=name, onesheet=onesheet, images=images, then=then)

  if type(then).__name__ == 'function':
    return then({'error': 'Bad filetype to insert images into. {0} must be of .docx or .xlsx'.format(name) })


def addImagesToMultipleSheets(name=testfile, images=testimages, then=None):
  addImages(name=name, images=images, onesheet=False, then=then)
  

def addImagesToOneSheet(name=testfile, images=testimages, then=None):
  addImages(name=name, images=images, onesheet=True, then=then)


if __name__ == "__main__":
  print('running as test')
  outfile = os.path.realpath(testfile).replace('inputs', 'outputs')
  if(os.path.exists(outfile)):
    os.remove(outfile)
  shutil.copy(testfile, outfile)

  addImagesToMultipleSheets(name=outfile, then=lambda x: print({'success': x}) )
  # createAll()
  print('success!')